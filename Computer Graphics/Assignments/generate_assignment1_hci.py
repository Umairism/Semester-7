from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = Path(__file__).with_name("Assignment 1 - Completed.docx")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_wireframe(doc, title, lines):
    doc.add_paragraph(title, style="List Bullet")
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["No Spacing"]
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        if index != len(lines) - 1:
            run.add_break()


document = Document()
section = document.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

normal_style = document.styles["Normal"]
normal_style.font.name = "Calibri"
normal_style.font.size = Pt(11)

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Assignment 1 - Human Computer Interaction (HCI)")
run.bold = True
run.font.size = Pt(16)

subtitle = document.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Course: HCI & Computer Graphics\nSemester: BSCS - 7\nCLO-1: Explain key principles, models, and theories of Human-Computer Interaction (HCI)")

info = document.add_table(rows=4, cols=2)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
info.style = "Table Grid"
labels = ["Student Name", "Roll Number", "Instructor", "Submission Date"]
for row, label in zip(info.rows, labels):
    set_cell_text(row.cells[0], label, bold=True)
    set_cell_text(row.cells[1], "____________________________")

document.add_paragraph()

heading = document.add_paragraph()
heading.add_run("Part 1 - Interface Evaluation").bold = True

document.add_paragraph(
    "Selected application: Google Classroom. The evaluation below uses Jakob Nielsen's heuristic principles to identify common usability problems in the interface."
)

table = document.add_table(rows=1, cols=4)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Usability Problem", "Issue Explanation", "Violated Heuristic", "Suggested Improvement"]
for cell, text in zip(table.rows[0].cells, headers):
    set_cell_text(cell, text, bold=True)
    shade_cell(cell, "D9EAF7")

issues = [
    (
        "Assignments are spread across Stream, Classwork, and Calendar.",
        "Students often need to search in multiple places to find tasks, due dates, and learning materials. This increases cognitive load and makes it harder for first-time users to understand where important academic work is located.",
        "Recognition rather than recall",
        "Add a single unified Assignments view with filters for pending, submitted, and overdue work so students can find tasks without remembering which section contains them.",
    ),
    (
        "Submission status is not always prominent enough.",
        "A student may still feel uncertain about whether work has been successfully turned in, especially when switching between class pages and assignment pages. Important status information is visible, but not persistent across all related screens.",
        "Visibility of system status",
        "Show a persistent status badge such as Submitted, Missing, or Late on course cards and inside the assignment list, not only on the assignment detail page.",
    ),
    (
        "The home view becomes visually cluttered when many classes are active.",
        "Course cards, announcements, icons, and due items compete for attention. When a student is enrolled in many courses, the interface becomes harder to scan and the most important actions do not stand out clearly.",
        "Aesthetic and minimalist design",
        "Reduce visual noise by letting users pin key classes, collapse less important announcements, and highlight only the highest-priority upcoming tasks on the main dashboard.",
    ),
    (
        "Navigation labels are not always aligned with student mental models.",
        "New users may not immediately understand the difference between Stream and Classwork. The system uses platform-specific terms instead of labels that directly reflect student goals such as Assignments, Notes, or Announcements.",
        "Match between system and the real world",
        "Rename or supplement navigation labels with clearer student-facing terms and short helper text so users can predict where to go before clicking.",
    ),
]

for problem, explanation, heuristic, improvement in issues:
    row = table.add_row().cells
    set_cell_text(row[0], problem)
    set_cell_text(row[1], explanation)
    set_cell_text(row[2], heuristic)
    set_cell_text(row[3], improvement)

document.add_paragraph()
document.add_paragraph(
    "Summary: Google Classroom is effective for basic academic communication, but its task discovery, feedback visibility, and information hierarchy can be improved to reduce confusion and support faster student decision-making."
)

document.add_paragraph()
part2 = document.add_paragraph()
part2.add_run("Part 2 - Interface Design Activity").bold = True

document.add_paragraph("Chosen system: Student Portal")
document.add_paragraph(
    "Prototype type: Low-fidelity paper prototype. The design focuses on layout, navigation flow, and task structure rather than detailed colors, icons, or final visual polish."
)

document.add_paragraph(
    "This prototype follows user-centered design principles in the following ways:",
)
ucd_points = [
    "It prioritizes common student tasks such as login, checking notices, and course registration.",
    "It uses familiar labels like Dashboard, Courses, Fees, and Timetable to match user expectations.",
    "It keeps the navigation consistent across screens so users do not need to relearn the interface.",
    "It reduces memory load by showing clear buttons, visible status messages, and a simple top-to-bottom task flow.",
    "It supports feedback by showing confirmation messages after actions such as adding or submitting course selections.",
]
for point in ucd_points:
    document.add_paragraph(point, style="List Bullet")

document.add_paragraph()
document.add_paragraph("Reference wireframes for hand-drawn prototype:")

add_wireframe(
    document,
    "Login screen",
    [
        "+----------------------------------+",
        "|         STUDENT PORTAL          |",
        "|----------------------------------|",
        "| Student ID: [______________]     |",
        "| Password  : [______________]     |",
        "| [ Sign In ]   [ Forgot Password ]|",
        "| Help Desk: support@portal.edu    |",
        "+----------------------------------+",
    ],
)

add_wireframe(
    document,
    "Main menu / dashboard",
    [
        "+----------------------------------+",
        "| STUDENT PORTAL   Welcome, Ali    |",
        "|----------------------------------|",
        "| [ Dashboard ] [ Courses ] [ Fees]|",
        "| [ Timetable ] [ Notices ]        |",
        "|----------------------------------|",
        "| Upcoming Classes                 |",
        "| - HCI & Computer Graphics        |",
        "| - Compiler Construction          |",
        "|----------------------------------|",
        "| Notices: Mid-term schedule live  |",
        "+----------------------------------+",
    ],
)

add_wireframe(
    document,
    "Functional screen: Course registration",
    [
        "+----------------------------------+",
        "| COURSE REGISTRATION              |",
        "|----------------------------------|",
        "| Search Course: [____________]    |",
        "| Available Courses                |",
        "| [ ] CS471 HCI                    |",
        "| [ ] CS473 Information Security   |",
        "| [ ] CS475 Compiler Construction  |",
        "|----------------------------------|",
        "| Selected: 0    Credit Hours: 0   |",
        "| [ Add ]   [ Remove ] [ Submit ]  |",
        "+----------------------------------+",
    ],
)

document.add_paragraph()
document.add_paragraph(
    "Instructions for final submission: Draw these three screens by hand on paper, take clear pictures, and insert those pictures below the relevant headings in the final document if required by your instructor."
)

photo_headings = [
    "Insert hand-drawn login screen photo here.",
    "Insert hand-drawn dashboard photo here.",
    "Insert hand-drawn course registration photo here.",
]
for text in photo_headings:
    p = document.add_paragraph()
    p.add_run(text).italic = True
    document.add_paragraph("\n")

document.add_paragraph()
conclusion = document.add_paragraph()
conclusion.add_run("Conclusion:").bold = True
conclusion.add_run(
    " The heuristic evaluation shows that even widely used systems can have usability gaps. The prototype demonstrates how a simple student portal can be designed around clarity, consistency, and common student needs using a low-fidelity user-centered approach."
)

document.save(OUTPUT_PATH)
print(f"Created {OUTPUT_PATH}")